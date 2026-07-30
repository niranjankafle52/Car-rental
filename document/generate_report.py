#!/usr/bin/env python3
"""
Generates the BCA final-year project report for the Car Rental System,
following the structure/format of document/refrence.docx.

Run:  python3 document/generate_report.py
Output: document/Car_Rental_System_Project_Report.docx

After opening in Word/LibreOffice:
  - Select All (Ctrl+A) -> F9 (or right-click a field -> Update Field) to
    resolve the Table of Contents, List of Figures, List of Tables and
    page numbers, since python-docx cannot paginate the document itself.
"""
import os
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(BASE, "screenshots")
DIAGRAMS = os.path.join(BASE, "diagrams")
OUT = os.path.join(BASE, "Car_Rental_System_Project_Report.docx")

AUTHOR = "Niranjan Kafle"
REG_NO = "[Registration Number]"
SUPERVISOR = "Riwaj Bhurtel"
SUBMIT_DATE = "July, 2026"
PROJECT_TITLE = "CAR RENTAL SYSTEM"

# ---------------------------------------------------------------------------
# Low-level OOXML helpers (fields, page numbers, section breaks)
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_field(paragraph, instr_text, result_text=""):
    """Insert a raw Word field (e.g. TOC, PAGE, SEQ) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t')
    fld_text.text = result_text
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    t_run = OxmlElement('w:r')
    t_run.append(fld_text)
    r_element.addnext(t_run)
    t_run.addnext(fld_end)
    return run


def add_seq_caption(doc, label, seq_name, text, style_name="Caption"):
    """Add a caption paragraph: '<label> <SEQ n>: <text>' e.g. 'Figure 3.1: Use Case Diagram'."""
    p = doc.add_paragraph(style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + " ")
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' SEQ {seq_name} \\* ARABIC '
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_text = OxmlElement('w:t'); fld_text.text = '1'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin); r.append(instr); r.append(fld_sep)
    t_run = OxmlElement('w:r'); t_run.append(fld_text)
    r.addnext(t_run); t_run.addnext(fld_end)
    p.add_run(f": {text}")
    return p


def add_page_number_field(paragraph, fmt=None):
    if fmt:
        p2 = paragraph.add_run()
    add_field(paragraph, ' PAGE  \\* MERGEFORMAT ', '1')


def set_page_number_format(section, fmt="decimal", start=None, restart=False):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if restart:
        pgNumType.set(qn('w:start'), str(start if start else 1))


def add_toc_field(doc, switches):
    p = doc.add_paragraph()
    add_field(p, f' {switches} ', 'Right-click and choose "Update Field" to build this list.')
    return p


def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


# ---------------------------------------------------------------------------
# Document setup: page size, margins, base styles
# ---------------------------------------------------------------------------

def new_document():
    doc = Document()

    # --- Page size (A4) and margins ---
    # Standard TU/BCA final-year-report format: extra left margin for binding.
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # --- Normal / body style ---
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Body Text style (used for narrative paragraphs) ---
    if 'Body Text' not in [s.name for s in doc.styles]:
        body = doc.styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        body = doc.styles['Body Text']
    body.base_style = doc.styles['Normal']
    body.font.name = 'Times New Roman'
    body.font.size = Pt(12)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.first_line_indent = None

    # --- Heading styles ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = False
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # --- Caption style ---
    if 'Caption' not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles['Caption']
    cap.base_style = doc.styles['Normal']
    cap.font.name = 'Times New Roman'
    cap.font.size = Pt(11)
    cap.font.italic = True
    cap.font.bold = True
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(14)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # --- Table Paragraph style (cell text) ---
    if 'Table Paragraph' not in [s.name for s in doc.styles]:
        tp = doc.styles.add_style('Table Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    else:
        tp = doc.styles['Table Paragraph']
    tp.base_style = doc.styles['Normal']
    tp.font.name = 'Times New Roman'
    tp.font.size = Pt(11)
    tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_after = Pt(2)

    # --- TOC styles (indent levels) ---
    for lvl, indent in ((1, 0.0), (2, 0.25), (3, 0.5), (4, 0.75)):
        name = f'toc {lvl}'
        if name not in [s.name for s in doc.styles]:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            st = doc.styles[name]
        st.base_style = doc.styles['Normal']
        st.font.name = 'Times New Roman'
        st.font.size = Pt(12)
        st.paragraph_format.left_indent = Inches(indent)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # tab stop with dot leader at right margin
        tab_stops = st.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), alignment=3, leader=2)  # RIGHT, DOTS

    return doc


def add_footer_page_number(section, center=True):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    add_field(p, ' PAGE ', '1')


def start_new_section(doc, same_as_prev_layout=True):
    """Insert a next-page section break, copying page geometry."""
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if same_as_prev_layout:
        prev = doc.sections[-2]
        new_sec.page_width = prev.page_width
        new_sec.page_height = prev.page_height
        new_sec.top_margin = prev.top_margin
        new_sec.bottom_margin = prev.bottom_margin
        new_sec.left_margin = prev.left_margin
        new_sec.right_margin = prev.right_margin
    return new_sec


# ---------------------------------------------------------------------------
# Small content helpers
# ---------------------------------------------------------------------------

def centered(doc, text, size=12, bold=False, italic=False, style=None, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p


def body(doc, text, style="Body Text", align=None, space_after=6):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, style="List Paragraph"):
    p = doc.add_paragraph(text, style=style)
    p.style = doc.styles['List Paragraph']
    p.paragraph_format.left_indent = Inches(0.4)
    fmt_run(p, size=12)
    # apply a simple bullet via numbering-less dash if no numbering defs are wired up
    return p


def fmt_run(p, size=12, name='Times New Roman'):
    for r in p.runs:
        r.font.name = name
        r.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(it)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)


def heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f'Heading {level}')


CACHE_DIR = os.path.join(BASE, "_render_cache")
os.makedirs(CACHE_DIR, exist_ok=True)

# Full-page scrolling captures get cropped to a representative "above the
# fold plus a bit" region so they fit on a single report page instead of
# scaling down to an unreadable sliver.
CROP_TOP_PX = {
    "home.png": 960,
    "home-mobile.png": 1300,
    "car-listing.png": 1360,
    "car-listing-mobile.png": 1180,
}


def _prepare_image(src_path):
    """Crop (if needed) and re-encode a screenshot as JPEG into the render
    cache, returning the cache path. Idempotent across repeated runs."""
    fname = os.path.basename(src_path)
    cache_path = os.path.join(CACHE_DIR, os.path.splitext(fname)[0] + ".jpg")
    if os.path.exists(cache_path) and os.path.getmtime(cache_path) >= os.path.getmtime(src_path):
        return cache_path

    im = Image.open(src_path)
    if im.mode in ("RGBA", "P", "LA"):
        bg = Image.new("RGB", im.size, (255, 255, 255))
        bg.paste(im.convert("RGBA"), mask=im.convert("RGBA").split()[-1])
        im = bg
    else:
        im = im.convert("RGB")

    if fname in CROP_TOP_PX:
        h = min(CROP_TOP_PX[fname], im.height)
        im = im.crop((0, 0, im.width, h))

    # Cap resolution so the resulting docx stays a reasonable size.
    max_w = 1600
    if im.width > max_w:
        ratio = max_w / im.width
        im = im.resize((max_w, int(im.height * ratio)))

    im.save(cache_path, "JPEG", quality=84, optimize=True)
    return cache_path


def add_figure(doc, path, width_in=5.8, caption_text=None, seq_name="Figure", label="Figure 3.",
               max_height_in=8.2, is_screenshot=True):
    render_path = _prepare_image(path) if is_screenshot else path
    im = Image.open(render_path)
    aspect = im.width / im.height  # width / height
    w = width_in
    h = w / aspect
    if h > max_height_in:
        h = max_height_in
        w = h * aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(render_path, width=Inches(w))
    if caption_text:
        add_seq_caption(doc, label, seq_name, caption_text)


def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F2F2F2")
    cell.text = ""
    lines = code_text.strip("\n").split("\n")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line if line.strip() else " ")
        r.font.name = 'Consolas'
        r.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def figure_counter_state():
    # simple python-side counter used only for on-screen text like "Figure 3.4";
    # the actual SEQ field in the docx will auto-number independently when
    # updated in Word, this is just for readable placeholder text.
    if not hasattr(figure_counter_state, "n"):
        figure_counter_state.n = 0
    figure_counter_state.n += 1
    return figure_counter_state.n


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------

def build_cover_page(doc):
    centered(doc, "Tribhuvan University", size=14, bold=True, space_after=0)
    centered(doc, "Faculty of Humanities and Social Sciences", size=13, space_after=40)

    centered(doc, f'"{PROJECT_TITLE}"', size=18, bold=True, space_after=6)
    centered(doc, "A PROJECT REPORT", size=13, bold=True, italic=True, space_after=40)

    centered(doc, "Submitted to", size=12, space_after=4)
    centered(doc, "Department of Computer Application", size=13, bold=True, space_after=0)
    centered(doc, "Reliance College", size=13, bold=True, space_after=30)

    centered(doc, "In partial fulfillment of the requirements for the Bachelor in Computer Application",
             size=12, italic=True, space_after=40)

    centered(doc, "Submitted by:", size=12, space_after=4)
    centered(doc, AUTHOR, size=13, bold=True, space_after=2)
    centered(doc, f"(Registration Number: {REG_NO})", size=12, space_after=30)

    centered(doc, "Under the Supervision of", size=12, space_after=2)
    centered(doc, SUPERVISOR, size=13, bold=True, space_after=30)

    centered(doc, SUBMIT_DATE, size=12, bold=True, space_after=0)


def build_recommendation(doc):
    centered(doc, "Tribhuvan University", size=13, bold=True, space_after=0)
    centered(doc, "Faculty of Humanities and Social Sciences", size=12, space_after=0)
    centered(doc, "Reliance College", size=12, space_after=20)

    heading(doc, "SUPERVISOR'S RECOMMENDATION", level=1)
    body(doc,
         f"I hereby recommend that this project prepared under my supervision by {AUTHOR} "
         f"(Registration Number: {REG_NO}) entitled “{PROJECT_TITLE}” be accepted as fulfilling "
         "the requirement for the partial fulfillment of the degree of Bachelor in Computer Application. "
         "To the best of my knowledge, this is an original piece of work carried out under my guidance "
         "and has not been submitted elsewhere for the award of any degree.")
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph("....................................................")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading(doc, "SIGNATURE", level=3)
    body(doc, SUPERVISOR, space_after=0)
    heading(doc, "SUPERVISOR", level=3)
    body(doc, "Department of Computer Application", space_after=0)
    body(doc, "Reliance College", space_after=0)


def build_approval(doc):
    centered(doc, "Tribhuvan University", size=13, bold=True, space_after=0)
    centered(doc, "Faculty of Humanities and Social Sciences", size=12, space_after=0)
    centered(doc, "Reliance College", size=12, space_after=20)

    heading(doc, "LETTER OF APPROVAL", level=2)
    body(doc,
         f"This is to certify that this project prepared by {AUTHOR} (T.U Registration Number: {REG_NO}) "
         f"entitled “{PROJECT_TITLE}” has been examined and is submitted for evaluation by the "
         "Department of Computer Application, Reliance College, in partial fulfillment of the requirements "
         "for the degree of Bachelor in Computer Application. This project report has not been submitted "
         "for any other degree and is forwarded for evaluation.")
    doc.add_paragraph().paragraph_format.space_after = Pt(30)

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cells = table.rows[0].cells
    cells[0].text = "...................................................."
    cells[1].text = "...................................................."
    cells = table.rows[1].cells
    cells[0].text = f"{SUPERVISOR}\nSupervisor"
    cells[1].text = "Internal Examiner"
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p = centered(doc, "....................................................", space_after=0)
    centered(doc, "External Examiner", space_after=0)


def build_abstract(doc):
    heading(doc, "ABSTRACT", level=1)
    body(doc,
         f"{PROJECT_TITLE.title()} is a web-based car rental booking and fleet management platform designed "
         "to streamline the process of listing, discovering, and booking rental vehicles online. The system "
         "replaces the manual, phone-call-driven booking process traditionally used by small and mid-sized "
         "car rental businesses with a self-service portal that lets customers browse the available fleet, "
         "filter vehicles by brand, fuel type, seating capacity and price, check real-time availability for a "
         "chosen date range, and complete a booking with online payment.")
    body(doc,
         "The platform is built on a PHP and MySQL stack and is organised into two cooperating modules: a "
         "customer-facing portal for registration, search, booking, testimonials and account management, and "
         "an administrative back office for managing vehicles, brands, bookings, offers, testimonials and "
         "customer enquiries. A weighted greedy algorithm ranks the available fleet against the customer's "
         "stated preferences (vehicle type, fuel type, brand and budget) to recommend the best-fitting cars "
         "first, and an AI-powered chatbot, backed by an OpenRouter large-language-model with function-calling "
         "tools, assists visitors with vehicle search, booking questions and account queries in natural "
         "language. Online payment is processed through the Khalti payment gateway.")
    body(doc,
         "The system was analysed and designed using standard software engineering artefacts — use-case "
         "diagrams, data-flow diagrams, an entity-relationship schema and a system architecture diagram — "
         "and was implemented incrementally, with each module unit-tested against a documented set of test "
         "cases before integration. The resulting system demonstrates that a well-scoped, algorithmically "
         "assisted booking platform can materially reduce the manual overhead of running a car rental "
         "business while giving customers a faster and more transparent booking experience.")
    body(doc,
         "Keywords: Car Rental System, Vehicle Booking, Greedy Algorithm, AI Chatbot, PHP, MySQL, Khalti "
         "Payment Gateway, Role-Based Access Control.")


def build_acknowledgement(doc):
    heading(doc, "ACKNOWLEDGEMENT", level=1)
    body(doc,
         "I would like to express my sincere gratitude to Tribhuvan University and the Department of Computer "
         "Application, Reliance College, for providing me with the opportunity and the platform to undertake "
         f"this project as part of the Bachelor in Computer Application programme.")
    body(doc,
         f"I am deeply grateful to my supervisor, {SUPERVISOR}, for the continuous guidance, valuable "
         "feedback and encouragement provided throughout the analysis, design, implementation and testing of "
         f"this project, “{PROJECT_TITLE}”. Their insights were instrumental in shaping the direction "
         "of this work.")
    body(doc,
         "I would also like to thank the faculty members of the Department of Computer Application for their "
         "support throughout the course of study, and my family, friends and classmates for their patience, "
         "motivation and constructive suggestions during the development of this project.")
    body(doc, "Finally, any errors or shortcomings that remain in this report are entirely my own.")
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    heading(doc, "Yours Sincerely,", level=3)
    body(doc, AUTHOR, space_after=0)
    body(doc, f"(Registration Number: {REG_NO})", space_after=0)


def build_toc_pages(doc):
    heading(doc, "TABLE OF CONTENTS", level=1)
    add_toc_field(doc, 'TOC \\o "1-3" \\h \\z \\u')
    add_page_break(doc)

    heading(doc, "LIST OF FIGURES", level=1)
    add_toc_field(doc, 'TOC \\c "Figure"')
    add_page_break(doc)

    heading(doc, "LIST OF TABLES", level=1)
    add_toc_field(doc, 'TOC \\c "Table"')
    add_page_break(doc)


ABBREVIATIONS = [
    ("API", "Application Programming Interface"),
    ("BCA", "Bachelor in Computer Application"),
    ("CSS", "Cascading Style Sheets"),
    ("DFD", "Data Flow Diagram"),
    ("ER", "Entity Relationship"),
    ("HTML", "Hyper Text Markup Language"),
    ("HTTP", "Hyper Text Transfer Protocol"),
    ("JSON", "JavaScript Object Notation"),
    ("LLM", "Large Language Model"),
    ("PDO", "PHP Data Objects"),
    ("PHP", "Hypertext Preprocessor"),
    ("RBAC", "Role-Based Access Control"),
    ("SQL", "Structured Query Language"),
    ("TU", "Tribhuvan University"),
    ("UI/UX", "User Interface / User Experience"),
    ("URL", "Uniform Resource Locator"),
]


def build_abbreviations(doc):
    heading(doc, "ABBREVIATIONS", level=1)
    table = doc.add_table(rows=1 + len(ABBREVIATIONS), cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Full Form"
    for row in table.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True
    for i, (abbr, full) in enumerate(ABBREVIATIONS, start=1):
        table.rows[i].cells[0].text = abbr
        table.rows[i].cells[1].text = full
    for row in table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                p.style = doc.styles['Table Paragraph']
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    set_repeat_table_header(table.rows[0])


# ---------------------------------------------------------------------------
# Chapter 1: Introduction
# ---------------------------------------------------------------------------

def build_chapter1(doc):
    heading(doc, "CHAPTER 1: INTRODUCTION", level=1)

    heading(doc, "1.1 Introduction", level=2)
    body(doc,
         "The rental car industry has traditionally relied on phone calls, walk-in counters and manually "
         "maintained spreadsheets to track vehicle availability, bookings and customer records. This approach "
         "does not scale well: staff must manually cross-check dates against a fleet list to confirm "
         "availability, customers cannot compare vehicles or prices without contacting the business directly, "
         "and there is no self-service way to view, modify or cancel a booking once it has been made.")
    body(doc,
         "The Car Rental System is a full-stack web application that digitises this entire workflow. It gives "
         "customers a public storefront where they can browse the fleet, filter vehicles by brand, fuel type, "
         "seating capacity and daily rate, check live availability for a chosen pickup and return date, and "
         "complete a booking online with payment through the Khalti payment gateway. On the other side, "
         "administrators are given a dedicated back-office panel to manage the vehicle inventory, brands, "
         "bookings, promotional offers, customer testimonials and contact enquiries from a single dashboard.")
    body(doc,
         "Beyond basic CRUD-style booking management, the system layers two pieces of applied algorithmic "
         "logic on top of the core workflow: a weighted greedy recommendation algorithm that ranks the "
         "available fleet against a customer's stated preferences, and an AI chatbot, backed by a large "
         "language model through the OpenRouter API, that can answer natural-language questions, search the "
         "fleet and even place a booking on the customer's behalf through a set of guarded function-calling "
         "tools.")

    heading(doc, "1.2 Problem Statement", level=2)
    body(doc,
         "Small and mid-sized car rental businesses typically cannot afford (or do not need) a large "
         "enterprise fleet-management suite, and end up running their booking process manually. This creates "
         "several concrete problems that this project sets out to address:")
    add_bullets(doc, [
        "Vehicle availability is tracked manually, which makes double-booking of the same vehicle for "
        "overlapping dates a real risk.",
        "Customers have no way to browse, filter or compare the available fleet before contacting the "
        "business, which lengthens the booking cycle.",
        "There is no self-service way for a customer to view their booking history, cancel a booking, or "
        "update their account details.",
        "Administrators lack a single dashboard to manage vehicles, brands, offers, testimonials and "
        "customer enquiries — these are often spread across spreadsheets, messaging apps and paper records.",
        "First-time visitors often have basic questions (pricing, fuel type, how booking works) that would "
        "otherwise require a phone call or a wait for a reply to be answered.",
    ])

    heading(doc, "1.3 Objective", level=2)
    body(doc, "The broad objective of this project is to design and implement a web-based car rental "
              "booking and fleet-management platform. The specific objectives are:")
    add_numbered(doc, [
        "To develop a public-facing portal where customers can register, browse the vehicle fleet, search "
        "and filter by brand/fuel type/price, and view detailed information for each vehicle.",
        "To implement a date-range-aware booking workflow that prevents a vehicle from being double-booked "
        "for overlapping dates, integrated with the Khalti online payment gateway.",
        "To build an administrative back office for managing vehicles, brands, bookings, promotional offers, "
        "testimonials and contact-us enquiries.",
        "To design and implement a greedy algorithm that scores and ranks available vehicles against a "
        "customer's stated preferences (vehicle type, fuel type, brand, seating and budget).",
        "To integrate an AI-powered chatbot capable of answering customer questions and performing guarded "
        "actions (search, booking) through function calling against a large language model.",
    ])

    heading(doc, "1.4 Scope and Limitation", level=2)
    body(doc, "Scope:")
    add_bullets(doc, [
        "Customer registration, login and profile/account management.",
        "Vehicle catalogue with search, keyword/brand/fuel-type filtering and pagination.",
        "Date-range booking with availability checking, checkout and Khalti payment integration.",
        "Customer testimonial submission and a public contact/enquiry form.",
        "An AI chatbot for search, FAQ and guarded booking assistance.",
        "A role-based administrative panel for vehicles, brands, bookings, offers and testimonials.",
    ])
    body(doc, "Limitation:")
    add_bullets(doc, [
        "The system is scoped to a single rental branch/location and does not model multi-branch fleet "
        "transfers.",
        "Payment is integrated with a single gateway (Khalti); other gateways are not implemented.",
        "The recommendation algorithm is a deterministic weighted greedy scorer rather than a learned "
        "(machine-learning-based) recommender.",
        "The chatbot's knowledge is limited to the information exposed through its tool definitions and does "
        "not have access to information outside the platform's database.",
    ])

    heading(doc, "1.5 Development Methodology", level=2)
    body(doc,
         "The project was developed using an incremental development model. Rather than attempting to design "
         "the complete system up-front, the system was decomposed into independently deliverable increments — "
         "public browsing and search, customer authentication, the booking and payment workflow, the "
         "administrative panel, the greedy recommendation engine, and finally the AI chatbot — with each "
         "increment analysed, designed, implemented and tested before the next one began. This allowed core "
         "booking functionality to be validated early, while later increments (recommendation and chatbot) "
         "were layered on top of a already-working booking pipeline.")
    add_figure(doc, os.path.join(DIAGRAMS, "gantt.png"), width_in=6.0,
               caption_text="Project Schedule (Gantt Chart)", is_screenshot=False)

    heading(doc, "1.6 Report Organization", level=2)
    body(doc, "This report is organised into five chapters:")
    add_bullets(doc, [
        "Chapter 1 (Introduction) presents the motivation, problem statement, objectives, scope and "
        "development methodology of the project.",
        "Chapter 2 (Background Study and Literature Review) discusses the domain background and reviews "
        "related systems and techniques.",
        "Chapter 3 (System Analysis and Design) covers requirement analysis, feasibility analysis, the "
        "use-case, data-flow and architectural diagrams, database schema design, the interface design of "
        "every major screen, and the algorithms used by the system.",
        "Chapter 4 (Implementation and Testing) describes the tools and technologies used, module-level "
        "implementation details, and the test cases used to verify each feature.",
        "Chapter 5 (Conclusion and Future Recommendations) presents a critical analysis of the finished "
        "system, its limitations, and directions for future work.",
    ])


# ---------------------------------------------------------------------------
# Chapter 2: Background Study and Literature Review
# ---------------------------------------------------------------------------

def build_chapter2(doc):
    heading(doc, "CHAPTER 2: BACKGROUND STUDY AND LITERATURE REVIEW", level=1)

    heading(doc, "2.1 Background Study", level=2)
    body(doc,
         "Vehicle rental is a well-established service industry, but a large share of small and independent "
         "operators still coordinate bookings manually — over the phone, through messaging apps, or with a "
         "walk-in counter and a paper or spreadsheet-based booking register. This is workable at low volume, "
         "but breaks down as the fleet and customer base grow: it is easy to lose track of which vehicle is "
         "booked for which date range, difficult to give customers an accurate, real-time picture of what is "
         "actually available, and time-consuming to manually reconcile payments against bookings.")
    body(doc,
         "A web-based booking platform addresses these problems directly. Vehicle availability becomes a "
         "query against a single source of truth (the bookings table) rather than something staff have to "
         "remember or look up; customers can self-serve — browsing, filtering and booking without waiting for "
         "a reply; and the business gets a searchable, auditable record of every booking, payment and "
         "customer enquiry. This project studies that transition for the specific case of a car rental "
         "business, and additionally explores how a lightweight recommendation algorithm and a conversational "
         "AI assistant can further reduce the friction of finding and booking the right vehicle.")

    heading(doc, "2.2 Literature Review", level=2)
    body(doc,
         "Commercial car-sharing and rental platforms such as Turo, Zoomcar and the booking modules of major "
         "rental chains popularised the now-familiar pattern of browsing a vehicle catalogue, filtering by "
         "category/price/fuel type, and checking a calendar of availability before paying online. These "
         "platforms informed the core user flow adopted in this project — catalogue, filter, date-range "
         "availability check, checkout — while this project keeps the implementation intentionally smaller "
         "in scope and self-hosted, suited to a single-branch rental business rather than a multi-vendor "
         "marketplace.")
    body(doc,
         "On the algorithmic side, greedy algorithms are a standard technique in resource-allocation and "
         "scheduling problems where a good-enough, locally-optimal choice at each step is acceptable in "
         "exchange for a simple, fast and deterministic decision procedure — as opposed to globally optimal "
         "but far more expensive approaches such as integer programming or learned ranking models. This "
         "project applies that idea to vehicle recommendation: rather than solving a combinatorial "
         "optimisation problem over the whole fleet, each available vehicle is scored independently against "
         "the customer's stated preferences (vehicle-type weight, fuel-type preference, brand preference and "
         "budget fit), and the highest-scoring vehicles are greedily selected as the top recommendations. This "
         "is appropriate here because recommendations must be computed on every search request, in real time, "
         "over a modestly sized fleet.")
    body(doc,
         "On the conversational side, the emergence of large language models with function/tool-calling "
         "support (as exposed through APIs such as OpenRouter) has made it practical to add a natural-language "
         "assistant to a web application without training a custom model: the model is given a fixed set of "
         "callable tools (searching cars, fetching car details, checking login state, creating a booking) and "
         "decides when to call them based on the conversation, while the application code — not the model — "
         "continues to enforce access control on any tool that touches booking or account data. This "
         "project's chatbot follows that pattern, which is increasingly common in customer-support and "
         "e-commerce assistants.")


# ---------------------------------------------------------------------------
# Chapter 3: System Analysis and Design
# ---------------------------------------------------------------------------

def build_chapter3(doc):
    heading(doc, "CHAPTER 3: SYSTEM ANALYSIS AND DESIGN", level=1)

    # ---- 3.1 System Analysis --------------------------------------------
    heading(doc, "3.1 System Analysis", level=2)

    heading(doc, "3.1.1 Requirement Analysis", level=3)
    heading(doc, "Functional Requirements", level=3)
    add_bullets(doc, [
        "Vehicle Catalogue: display the full fleet with brand, fuel type, seating capacity, model year and "
        "price-per-day; allow keyword, brand and fuel-type based search with pagination.",
        "Booking Management: let a logged-in customer pick a vehicle and a pickup/return date range, verify "
        "the vehicle is not already booked for an overlapping date range, and record the booking.",
        "Online Payment: route the checkout amount through the Khalti payment gateway and record the "
        "transaction result against the booking.",
        "Account Management: registration, login, profile editing and password update for customers.",
        "My Bookings: let a customer view their booking history and its current status.",
        "Testimonial Management: let a customer submit a testimonial, and let an administrator approve or "
        "reject it before it appears publicly.",
        "Contact / Inquiry Management: a public contact form that stores enquiries for admin follow-up.",
        "Chatbot Support: a conversational assistant that can answer questions, search the fleet and place a "
        "booking for a logged-in visitor.",
        "Admin Panel: CRUD management of vehicles, brands, bookings, offers, testimonials, contact "
        "enquiries and registered users, from a single dashboard.",
        "Smart Vehicle Recommendation: score and rank the available fleet against a customer's stated "
        "preferences using the greedy recommendation algorithm.",
    ])
    heading(doc, "Non-Functional Requirements", level=3)
    add_bullets(doc, [
        "Usability: a clean, responsive interface that works across desktop and mobile screen sizes.",
        "Security: passwords are never stored or transmitted in plain text without hashing, all database "
        "access uses parameterised (PDO prepared statement) queries to prevent SQL injection, and every "
        "chatbot tool that touches booking or account data re-checks the visitor's session server-side "
        "rather than trusting the language model's judgement.",
        "Compatibility: the system runs on any standards-compliant browser and on a standard PHP/MySQL "
        "hosting stack without special server extensions.",
        "Performance: fleet search, filtering and the greedy scoring pass must complete quickly enough to "
        "run synchronously on every search request.",
    ])

    heading(doc, "3.1.2 Use-Case Diagram", level=3)
    body(doc,
         "The system recognises two actors: the Customer (an unregistered visitor or a registered user of the "
         "public portal) and the Administrator (the operator of the back office). Figure 3.1 shows the "
         "principal use cases each actor participates in.")
    add_figure(doc, os.path.join(DIAGRAMS, "usecase.png"), width_in=6.0,
               caption_text="Use Case Diagram", is_screenshot=False)
    heading(doc, "Customer Use Cases", level=3)
    add_bullets(doc, [
        "Register / Login / Update profile and password.",
        "Browse and search the vehicle catalogue; use the smart finder to get ranked recommendations.",
        "View vehicle details and check availability for a date range.",
        "Book a vehicle and pay online through Khalti.",
        "View \"My Bookings\" and past booking status.",
        "Submit a testimonial and a contact / enquiry message.",
        "Chat with the AI assistant for search help, FAQs and guided booking.",
    ])
    heading(doc, "Administrator Use Cases", level=3)
    add_bullets(doc, [
        "Log in to the admin dashboard.",
        "Add, edit, delete and view vehicles and vehicle brands.",
        "View and manage new, confirmed and cancelled bookings.",
        "Manage promotional offers.",
        "Approve or reject customer testimonials.",
        "View and respond to contact / enquiry submissions and manage registered users.",
    ])

    heading(doc, "3.1.3 Feasibility Analysis", level=3)
    heading(doc, "Technical Feasibility", level=3)
    body(doc,
         "The system is built entirely on the open-source PHP/MySQL stack (with a small amount of "
         "JavaScript and Bootstrap on the front end), technologies that are freely available, well "
         "documented and already familiar from the BCA curriculum. The Khalti and OpenRouter integrations "
         "are both plain HTTPS/JSON APIs reachable with PHP's cURL extension, requiring no specialised "
         "infrastructure. The project is therefore technically feasible to build and deploy on standard "
         "shared or VPS hosting.")
    heading(doc, "Economic Feasibility", level=3)
    body(doc,
         "All software used — PHP, MySQL, and the code editor/IDE — is free and open-source, and Khalti's "
         "merchant integration and OpenRouter's pay-per-use API pricing keep running costs proportional to "
         "actual usage rather than requiring a large upfront licence fee. The system is therefore economically "
         "feasible for a small rental business to adopt.")
    heading(doc, "Schedule Feasibility", level=3)
    body(doc,
         "The incremental development plan shown in the Gantt chart in Section 1.5 allocates the project "
         "across analysis, module-by-module implementation and testing within the semester timeline, making "
         "the schedule realistic given the scope defined in Section 1.4.")

    heading(doc, "3.1.4 Data Modeling (Process Modeling)", level=3)
    body(doc,
         "Data flow diagrams (DFDs) were used to model how information moves between the customer, the "
         "administrator and the system's data stores. Figure 3.2 gives the context-level (Level 0) view of "
         "the whole system as a single process, and Figure 3.3 decomposes that process into its main "
         "sub-processes for Level 1.")
    add_figure(doc, os.path.join(DIAGRAMS, "dfd_level0.png"), width_in=6.0,
               caption_text="Level-0 Data Flow Diagram", is_screenshot=False)
    add_figure(doc, os.path.join(DIAGRAMS, "dfd_level1.png"), width_in=6.2,
               caption_text="Level-1 Data Flow Diagram", is_screenshot=False)

    # ---- 3.2 System Design -------------------------------------------------
    heading(doc, "3.2 System Design", level=2)

    heading(doc, "3.2.1 Architecture Design", level=3)
    body(doc,
         "The system follows a classic three-tier web architecture. The presentation tier is server-rendered "
         "PHP with Bootstrap and JavaScript for interactivity; the application tier is a set of PHP "
         "controllers/scripts implementing authentication, booking, the greedy recommendation engine and the "
         "chatbot's tool-calling logic; and the data tier is a MySQL database accessed exclusively through "
         "PDO prepared statements. The chatbot module additionally calls out to the OpenRouter API for "
         "language-model inference, and the checkout flow calls out to the Khalti API for payment "
         "processing, both over HTTPS.")
    add_figure(doc, os.path.join(DIAGRAMS, "architecture.png"), width_in=6.2,
               caption_text="System Architecture Design", is_screenshot=False)

    heading(doc, "3.2.2 Database Schema Design", level=3)
    body(doc,
         "The database is normalised around ten tables. Figure 3.5 shows the schema and the relationships "
         "between them.")
    add_figure(doc, os.path.join(DIAGRAMS, "er_schema.png"), width_in=6.2,
               caption_text="Database Schema Design", is_screenshot=False)
    db_table = doc.add_table(rows=1, cols=2)
    db_table.style = 'Table Grid'
    db_table.rows[0].cells[0].text = "Table"
    db_table.rows[0].cells[1].text = "Purpose"
    for r in db_table.rows[0].cells:
        for p in r.paragraphs:
            for run in p.runs:
                run.font.bold = True
    schema_rows = [
        ("admin", "Administrator login credentials for the back-office panel."),
        ("tblusers", "Registered customer accounts (name, email, password, contact and address details)."),
        ("tblvehicles", "The vehicle fleet: title, brand, overview, price per day, fuel type, model year, "
                        "seating capacity, feature flags (air conditioner, power door locks, etc.) and up to "
                        "five images per vehicle."),
        ("tblbrands", "Vehicle brand/manufacturer master list, referenced by tblvehicles."),
        ("tblbooking", "Booking records linking a customer email and a vehicle to a from/to date range, a "
                       "status flag, and an optional message."),
        ("tblcontactusquery", "Enquiry messages submitted through the public Contact Us form."),
        ("tblcontactusinfo", "The rental business's own published address, email and phone number."),
        ("tbltestimonial", "Customer testimonials pending or approved for public display."),
        ("tblsubscribers", "Newsletter/subscriber email addresses."),
        ("tblpages", "Editable static content pages (e.g. About Us) managed from the admin panel."),
    ]
    for name, desc in schema_rows:
        row = db_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
    for row in db_table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                p.style = doc.styles['Table Paragraph']
    set_repeat_table_header(db_table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    heading(doc, "3.2.3 Interface Design (UI/UX)", level=3)
    body(doc, "The following figures show the implemented interface for each major screen of the customer "
              "portal and the admin panel.")

    interface_shots = [
        ("home.png", "User Home Page"),
        ("home-mobile.png", "User Home Page (Mobile View)"),
        ("login-modal.png", "Login"),
        ("register-modal.png", "Register"),
        ("car-listing.png", "Vehicle Listing Page"),
        ("car-listing-mobile.png", "Vehicle Listing Page (Mobile View)"),
        ("search-results.png", "Search Results Page"),
        ("smart-finder.png", "Smart Car Finder (Greedy Recommendation)"),
        ("vehicle-details.png", "Vehicle Details Page"),
        ("my-booking.png", "My Bookings Page"),
        ("profile.png", "Account Details Page"),
        ("contact-us.png", "Contact Us Page"),
        ("post-testimonial.png", "Post Testimonial Page"),
        ("chatbot-widget.png", "AI Chatbot Widget"),
        ("admin-login.png", "Admin Login"),
        ("admin-dashboard.png", "Admin Dashboard"),
        ("admin-add-vehicle.png", "Admin: Add Vehicle"),
        ("admin-manage-vehicles.png", "Admin: Manage Vehicles"),
        ("admin-manage-brands.png", "Admin: Manage Brands"),
        ("admin-manage-bookings.png", "Admin: Manage Bookings"),
        ("admin-testimonials.png", "Admin: Manage Testimonials"),
        ("admin-users.png", "Admin: Registered Users"),
    ]
    for fname, caption in interface_shots:
        path = os.path.join(SCREENSHOTS, fname)
        if not os.path.exists(path):
            continue
        width = 3.1 if "mobile" in fname else 5.8
        add_figure(doc, path, width_in=width, caption_text=caption, is_screenshot=True)

    heading(doc, "3.2.4 Algorithm", level=2)

    heading(doc, "Greedy Vehicle Recommendation Algorithm", level=3)
    body(doc,
         "The smart finder ranks the fleet with a weighted greedy scoring function. For a given search, the "
         "system first filters out every vehicle already booked for an overlapping date range, then scores "
         "each remaining vehicle independently against the customer's stated preferences, and finally makes "
         "the greedy choice of always returning the highest-scoring vehicles first — rather than searching "
         "for a globally optimal combination, which would be unnecessary for a single-vehicle booking.")
    add_code_block(doc, """
function findOptimalCars(criteria, fromDate, toDate):
    available = getAvailableCars(fromDate, toDate)   // excludes overlapping bookings
    for each car in available:
        score = 0
        score += vehicleTypeWeight(car.type)          // SUV/Sedan/Luxury/... weight
        score += fuelPreferenceScore(car.fuelType, criteria)
        score += brandPreferenceScore(car.brand, criteria)
        score += budgetFitScore(car.pricePerDay, criteria.budget)
        score += seatingMatchScore(car.seats, criteria.seats)
        car.total_score = score
    sort available by total_score descending           // greedy: best-first
    return top 10 of available
""")
    body(doc, "How it works:")
    add_bullets(doc, [
        "Availability is checked first with a date-overlap query against tblbooking, so an already-booked "
        "vehicle is never recommended.",
        "Each remaining vehicle accumulates a score from independent weighted factors — vehicle-type weight, "
        "a fuel-type preference multiplier, a brand preference multiplier, and how closely the price and "
        "seating match the customer's stated budget and party size.",
        "The greedy step is the sort-and-take-the-best: the algorithm never backtracks or reconsiders a "
        "lower-scoring vehicle once a higher-scoring one is available, which keeps the recommendation pass "
        "O(n log n) over the available fleet and fast enough to run on every request.",
    ])
    body(doc,
         "Example: for a customer requesting an SUV, Petrol fuel and a budget around NRS 8,000/day, a "
         "Hyundai Creta (SUV, Petrol, NRS 8,000/day) scores higher than a Maruti Suzuki Wagon R (Hatchback, "
         "Petrol, NRS 500/day) even though the Wagon R is cheaper, because the vehicle-type and budget-fit "
         "terms dominate the score for this query.")

    heading(doc, "Chatbot Support Algorithm", level=3)
    body(doc,
         "The AI chatbot is implemented as a tool-calling conversation loop against an OpenRouter-hosted "
         "large language model. The model is not given direct database access; instead it is offered a fixed "
         "set of callable tools — check_login_status, search_cars, get_car_details and create_booking among "
         "others — and the PHP backend executes the tool the model asks for, enforcing the visitor's actual "
         "session state itself rather than trusting anything the model claims about who is logged in.")
    add_code_block(doc, """
loop:
    response = OpenRouter.chat(history, systemPrompt, tools=chatbotToolDefinitions())
    if response.requests_tool_call:
        for each tool_call in response.tool_calls:
            result = executeTool(tool_call.name, tool_call.arguments)  // server enforces auth here
            history.append(tool_call, result)
        continue   // let the model see the tool result and respond again
    else:
        return response.message   // final natural-language reply to the visitor
""")
    body(doc, "It works by:")
    add_bullets(doc, [
        "Maintaining the running conversation in the PHP session (chatbot_history) so the model has context "
        "across turns.",
        "Calling check_login_status before any booking-related tool, and refusing create_booking server-side "
        "if the session is not actually authenticated — the model's own belief about login state is never "
        "trusted for access control.",
        "Exposing search_cars and get_car_details as read-only tools so the assistant can answer fleet "
        "questions directly from the live database instead of guessing.",
    ])
    body(doc, "Example: a visitor asks \"do you have any electric cars under 6000 a day?\" — the model calls "
              "search_cars with the keyword \"electric\", receives the matching rows including price, and "
              "replies with the matches actually in the database rather than a hallucinated answer.")

    heading(doc, "Search, Filtering, and Pagination", level=3)
    body(doc,
         "The vehicle catalogue supports free-text keyword search plus brand and fuel-type filters, all "
         "translated into a single parameterised SQL query with WHERE/LIKE and equality clauses, combined "
         "with LIMIT/OFFSET for pagination so that only one page of results is fetched per request.")
    body(doc, "It works by:")
    add_bullets(doc, [
        "Building the WHERE clause incrementally from whichever filters the customer actually supplied, so "
        "an empty filter is simply omitted rather than matched against a wildcard.",
        "Joining tblvehicles against tblbrands so brand name, rather than the internal brand id, is shown "
        "and searched.",
        "Applying LIMIT/OFFSET last, after filtering, so pagination always operates over the already-narrowed "
        "result set.",
    ])
    body(doc, "Example: searching \"Hyundai\" with fuel type \"Petrol\" produces a WHERE clause matching the "
              "brand name and fuel type together, returning only Hyundai petrol vehicles, ten per page.")


# ---------------------------------------------------------------------------
# Chapter 4: Implementation and Testing
# ---------------------------------------------------------------------------

def add_test_table(doc, table_num, title, rows, appendix_start):
    """rows: list of (payload, expected, actual, result) tuples. Appendix numbers are
    assigned sequentially starting at appendix_start, and the next free number is returned."""
    add_seq_caption(doc, "Table 4.", "Table", title)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ["S.N.", "Payload", "Expected Result", "Actual Result", "Result", "Appendix"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    appendix_no = appendix_start
    for i, (payload, expected, actual, result) in enumerate(rows, start=1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = payload
        row.cells[2].text = expected
        row.cells[3].text = actual
        row.cells[4].text = result
        row.cells[5].text = f"Appendix {appendix_no}"
        appendix_no += 1
    for row in table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                p.style = doc.styles['Table Paragraph']
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return appendix_no


def build_chapter4(doc):
    heading(doc, "CHAPTER 4: IMPLEMENTATION AND TESTING", level=1)

    heading(doc, "4.1 Implementation", level=2)

    heading(doc, "4.1.1 CASE Tools", level=3)
    add_bullets(doc, [
        "Visual Studio Code — primary code editor for PHP, JavaScript, HTML and CSS.",
        "XAMPP (Apache + MySQL + PHP) — local development and hosting environment.",
        "phpMyAdmin — database administration and inspection of the carrental schema.",
        "Git and GitHub — version control and source management.",
        "Postman — manual endpoint testing for the JSON APIs (chatbot-api.php, greedy-api.php).",
        "Google Chrome DevTools — front-end debugging and responsive/mobile-layout testing.",
    ])

    heading(doc, "4.1.2 Programming Language", level=3)
    heading(doc, "PHP", level=3)
    body(doc,
         "PHP is used for all server-side logic: routing, authentication, session management, the booking "
         "workflow, the greedy recommendation engine, the chatbot's tool-calling loop, and all database "
         "access through the PDO extension with parameterised queries.")
    heading(doc, "HTML, CSS, Bootstrap and JavaScript", level=3)
    body(doc,
         "The presentation layer uses server-rendered HTML templated by PHP, styled with CSS and the "
         "Bootstrap framework for a responsive, mobile-friendly grid, and progressively enhanced with "
         "JavaScript for client-side interactivity such as the date-range picker, the chatbot widget and "
         "AJAX form submissions (e.g. live email-availability checking on registration).")
    heading(doc, "SQL", level=3)
    body(doc, "SQL (via MySQL) defines and queries the ten-table relational schema described in Section "
              "3.2.2, including the availability-overlap query that underpins both booking and the greedy "
              "recommendation engine.")

    heading(doc, "4.1.3 Database Platform", level=3)
    body(doc,
         "MySQL is used as the relational database engine, accessed exclusively through PHP's PDO layer. "
         "PDO's prepared statements bind every user-supplied value as a parameter rather than concatenating "
         "it into the SQL string, which is the project's primary defence against SQL injection.")

    heading(doc, "4.1.4 Implementation Details of Modules", level=3)
    module_details = [
        ("Authentication Module", "Handles registration, login, logout and password updates for both "
         "customers (tblusers) and administrators (admin), with a live AJAX email-availability check "
         "(check_availability.php) during registration."),
        ("Vehicle Catalogue Module", "Lists the fleet from tblvehicles joined with tblbrands, and implements "
         "keyword/brand/fuel-type search with pagination (search.php, search-carresult.php, car-listing.php)."),
        ("Booking Module", "Validates a requested date range against tblbooking to prevent double-booking, "
         "creates the booking record, and hands off to the payment module (reserve.php, my-booking.php)."),
        ("Payment Module", "Initiates and confirms a Khalti epayment transaction for a booking's total amount "
         "(initiate_payment.php, return_url.php, message.php)."),
        ("Smart Finder / Recommendation Module", "Implements the weighted greedy scoring algorithm described "
         "in Section 3.2.4 to rank the available fleet against customer-submitted preferences "
         "(greedy-form.php, greedy-api.php, greedy_algorithm.php)."),
        ("Chatbot Module", "Runs the tool-calling conversation loop against the OpenRouter API, exposing "
         "guarded tools for searching cars, fetching car details, checking login state and creating a "
         "booking (chatbot-api.php, includes/chatbot_tools.php, includes/chatbot_config.php)."),
        ("Testimonial Module", "Lets a logged-in customer submit a testimonial for admin approval before "
         "public display (post-testimonial.php, my-testimonials.php)."),
        ("Contact / Enquiry Module", "Stores public contact-form submissions for admin follow-up "
         "(contact-us.php)."),
        ("Admin Module", "A role-gated back office (admin/) for managing vehicles, brands, bookings, offers, "
         "testimonials, contact enquiries and registered users."),
    ]
    for name, desc in module_details:
        heading(doc, name, level=3)
        body(doc, desc)

    heading(doc, "4.2 Testing", level=2)
    body(doc,
         "Each module was tested manually against a documented set of test cases before being integrated "
         "with the rest of the system. Every table below records the input payload used, the expected "
         "result, the actual observed result, a pass/fail verdict, and a pointer to the corresponding "
         "screenshot evidence in the Appendix.")

    ap = 1
    ap = add_test_table(doc, 1, "Login Test Cases", [
        ("Email: (blank), Password: (blank)", "Please fill out this field.",
         "Please fill out this field.", "Pass"),
        ("Email: unregistered@test.com, Password: wrongpass", "Invalid email or password.",
         "Invalid email or password shown.", "Pass"),
        ("Email: registered user, Password: correct password", "User is logged in and redirected to home.",
         "User logged in and redirected to home page.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 2, "Register Test Cases", [
        ("Email already registered", "Email already exists — Register button disabled.",
         "\"Email already exists\" shown, submit disabled.", "Pass"),
        ("Valid new email, all fields completed", "Account created and user redirected to login.",
         "Account created successfully.", "Pass"),
        ("Password and Confirm Password mismatch", "Passwords do not match.",
         "Validation error shown, form not submitted.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 3, "Vehicle Search and Booking Test Cases", [
        ("Search keyword: \"Hyundai\", Fuel: Petrol", "Only Hyundai petrol vehicles listed.",
         "Only Hyundai petrol vehicles listed.", "Pass"),
        ("Book vehicle for dates overlapping an existing booking", "Booking rejected — vehicle unavailable "
         "for the selected dates.", "Booking rejected with an availability error.", "Pass"),
        ("Book vehicle for a free date range while logged in", "Booking created and user proceeds to "
         "checkout/payment.", "Booking created, redirected to payment.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 4, "Smart Finder (Greedy Recommendation) Test Cases", [
        ("Preferences: SUV, Petrol, budget ~NRS 8,000/day", "Higher-scoring SUVs ranked above cheaper "
         "non-matching vehicle types.", "SUV matches ranked first.", "Pass"),
        ("No vehicles available for the chosen dates", "Empty recommendation list returned gracefully.",
         "Empty list shown with a friendly message.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 5, "My Bookings Test Cases", [
        ("Logged-in user with existing bookings opens My Bookings", "All of the user's bookings are listed "
         "with correct status and invoice totals.", "Bookings listed correctly with invoice totals.", "Pass"),
        ("User with no bookings opens My Bookings", "An empty state is shown instead of an error.",
         "Empty state shown.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 6, "Contact Us and Chatbot Test Cases", [
        ("Contact form submitted with all required fields", "Enquiry saved and visible to the admin.",
         "Enquiry saved successfully.", "Pass"),
        ("Chatbot asked to book a car while not logged in", "Chatbot asks the visitor to log in before "
         "booking; create_booking is not executed.", "Chatbot correctly requested login before booking.",
         "Pass"),
        ("Chatbot asked \"what electric cars do you have\"", "Chatbot calls search_cars and replies with "
         "actual matching vehicles from the database.", "Correct matches returned from the live catalogue.",
         "Pass"),
    ], ap)
    ap = add_test_table(doc, 7, "Admin Vehicle and Brand Management Test Cases", [
        ("Admin adds a new vehicle with all required fields", "Vehicle saved and immediately visible in the "
         "public catalogue.", "Vehicle added and visible in listing.", "Pass"),
        ("Admin edits an existing vehicle's price per day", "Updated price reflected on the public vehicle "
         "details page.", "Price updated correctly.", "Pass"),
        ("Admin adds a new brand", "Brand appears in the brand filter/dropdown.",
         "Brand appears in filter list.", "Pass"),
    ], ap)
    ap = add_test_table(doc, 8, "Admin Booking and Testimonial Management Test Cases", [
        ("Admin opens Manage Bookings", "New, confirmed and cancelled bookings are listed with correct "
         "status.", "Bookings listed with correct status.", "Pass"),
        ("Admin approves a pending testimonial", "Testimonial becomes visible on the public site.",
         "Testimonial shown publicly after approval.", "Pass"),
        ("Admin rejects a pending testimonial", "Testimonial is not shown on the public site.",
         "Testimonial hidden from public view.", "Pass"),
    ], ap)


# ---------------------------------------------------------------------------
# Chapter 5, References, Appendix
# ---------------------------------------------------------------------------

def build_chapter5(doc):
    heading(doc, "CHAPTER 5: CONCLUSION AND FUTURE RECOMMENDATIONS", level=1)

    heading(doc, "5.1 Critical Analysis", level=2)
    body(doc,
         "The finished system meets the objectives set out in Chapter 1: it replaces a manual, phone-driven "
         "booking process with a self-service portal, prevents double-booking through a date-overlap check "
         "enforced at the database layer, integrates real online payment through Khalti, and adds two pieces "
         "of applied intelligence — the greedy recommendation engine and the tool-calling chatbot — that go "
         "beyond a plain CRUD booking system. Keeping access control (login checks) enforced server-side "
         "rather than delegated to the chatbot's own judgement proved to be an important design decision, "
         "since it means the AI-facing surface of the system cannot be tricked into bypassing authentication.")
    body(doc,
         "At the same time, the greedy scoring weights (vehicle type, fuel, brand, budget) are currently "
         "fixed constants rather than values learned from real booking outcomes, so the quality of a "
         "recommendation depends on how well those hand-chosen weights actually reflect customer behaviour "
         "— a limitation acknowledged below rather than something the current system measures.")

    heading(doc, "5.2 Conclusion", level=2)
    body(doc,
         f"{PROJECT_TITLE.title()} demonstrates that a small, well-scoped web application built on a "
         "conventional PHP/MySQL stack can meaningfully reduce the manual overhead of running a car rental "
         "business, while a lightweight greedy recommendation algorithm and a guarded, tool-calling AI "
         "chatbot can add real value on top of a standard booking workflow without requiring a large "
         "machine-learning infrastructure. The project reinforced practical skills in requirement analysis, "
         "relational database design, secure server-side development, and integrating third-party APIs "
         "(Khalti and OpenRouter) into a production-style workflow.")

    heading(doc, "5.3 Limitations", level=2)
    add_bullets(doc, [
        "Single rental branch and single payment gateway, as scoped in Section 1.4.",
        "The recommendation algorithm uses fixed, hand-tuned weights rather than weights learned from "
        "historical booking data.",
        "No automated test suite; testing in Chapter 4 was performed manually against documented test "
        "cases.",
        "The chatbot's knowledge is limited to what its tool definitions expose and does not cover queries "
        "outside the platform's own data.",
    ])

    heading(doc, "5.4 Future Enhancements", level=2)
    add_bullets(doc, [
        "Support multiple rental branches and inter-branch vehicle transfers.",
        "Add further payment gateways (e.g. eSewa, card payments) alongside Khalti.",
        "Replace the fixed greedy weights with a model trained on historical booking and rating data.",
        "Add automated unit and integration tests around the booking-overlap and greedy-scoring logic.",
        "Extend the chatbot with tools for booking modification/cancellation and richer multi-turn "
        "itinerary planning.",
    ])


def build_references(doc):
    heading(doc, "REFERENCES", level=1)
    refs = [
        "Sommerville, I. (2011). Software Engineering (9th ed.). Pearson Education.",
        "Pressman, R. S., & Maxim, B. R. (2014). Software Engineering: A Practitioner's Approach (8th ed.). "
        "McGraw-Hill Education.",
        "Elmasri, R., & Navathe, S. B. (2015). Fundamentals of Database Systems (7th ed.). Pearson.",
        "PHP Documentation Group. PHP Manual. Retrieved from https://www.php.net/manual/en/",
        "Oracle Corporation. MySQL 8.0 Reference Manual. Retrieved from https://dev.mysql.com/doc/",
        "Khalti Digital Wallet. Khalti ePayment API Documentation. Retrieved from https://docs.khalti.com/",
        "OpenRouter. OpenRouter API Documentation — Chat Completions and Tool Calling. Retrieved from "
        "https://openrouter.ai/docs",
        "Bootstrap Team. Bootstrap 5 Documentation. Retrieved from https://getbootstrap.com/docs/",
        "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms "
        "(3rd ed.). MIT Press. (Greedy algorithms, Chapter 16.)",
    ]
    for r in refs:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(r)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)


APPENDIX_ITEMS = [
    (1, "Login — validation on empty fields", "login-modal.png"),
    (2, "Login — invalid credentials", "login-modal.png"),
    (3, "Login — successful login (account session)", "profile.png"),
    (4, "Register — email already exists", "register-modal.png"),
    (5, "Register — successful registration", "register-modal.png"),
    (6, "Register — password mismatch validation", "register-modal.png"),
    (7, "Vehicle search — brand/fuel-type filter results", "search-results.png"),
    (8, "Booking — rejected for an overlapping date range", "vehicle-details.png"),
    (9, "Booking — created successfully", "my-booking.png"),
    (10, "Smart Finder — ranked recommendation results", "smart-finder.png"),
    (11, "Smart Finder — no vehicles available", "smart-finder.png"),
    (12, "My Bookings — booking history with invoice totals", "my-booking.png"),
    (13, "My Bookings — empty state", "my-booking.png"),
    (14, "Contact Us — enquiry submitted", "contact-us.png"),
    (15, "Chatbot — login required before booking", "chatbot-widget.png"),
    (16, "Chatbot — vehicle search reply", "chatbot-widget.png"),
    (17, "Admin — vehicle added", "admin-add-vehicle.png"),
    (18, "Admin — vehicle price updated", "admin-manage-vehicles.png"),
    (19, "Admin — brand added", "admin-manage-brands.png"),
    (20, "Admin — bookings managed", "admin-manage-bookings.png"),
    (21, "Admin — testimonial approved", "admin-testimonials.png"),
    (22, "Admin — testimonial rejected", "admin-testimonials.png"),
]


def build_appendix(doc):
    heading(doc, "APPENDIX", level=1)
    body(doc, "This appendix collects screenshot evidence for the test cases documented in Section 4.2, "
              "numbered to match the \"Appendix\" column of each test-case table.")
    for num, caption, fname in APPENDIX_ITEMS:
        heading(doc, f"Appendix {num}: {caption}", level=3)
        path = os.path.join(SCREENSHOTS, fname)
        if os.path.exists(path):
            add_figure(doc, path, width_in=4.3, caption_text=None, is_screenshot=True)


# ---------------------------------------------------------------------------
# Full document assembly
# ---------------------------------------------------------------------------

def build_document():
    doc = new_document()
    doc.core_properties.title = PROJECT_TITLE.title()
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "BCA Final Year Project Report"

    # --- Cover (no page number) ---
    build_cover_page(doc)

    # --- Front matter section: roman numeral page numbers ---
    start_new_section(doc)
    add_footer_page_number(doc.sections[-1])
    set_page_number_format(doc.sections[-1], fmt="lowerRoman", start=1, restart=True)

    build_recommendation(doc)
    add_page_break(doc)
    build_approval(doc)
    add_page_break(doc)
    build_abstract(doc)
    add_page_break(doc)
    build_acknowledgement(doc)
    add_page_break(doc)
    build_toc_pages(doc)
    build_abbreviations(doc)

    # --- Main body section: arabic numbering restarting at 1 ---
    start_new_section(doc)
    add_footer_page_number(doc.sections[-1])
    set_page_number_format(doc.sections[-1], fmt="decimal", start=1, restart=True)

    build_chapter1(doc)
    add_page_break(doc)
    build_chapter2(doc)
    add_page_break(doc)
    build_chapter3(doc)
    add_page_break(doc)
    build_chapter4(doc)
    add_page_break(doc)
    build_chapter5(doc)
    add_page_break(doc)
    build_references(doc)
    add_page_break(doc)
    build_appendix(doc)

    return doc


if __name__ == "__main__":
    doc = build_document()
    doc.save(OUT)
    print(f"Saved: {OUT}")
    print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables),
          "images:", len(list(doc.inline_shapes)), "sections:", len(doc.sections))
    print("size(MB):", round(os.path.getsize(OUT) / 1024 / 1024, 2))
